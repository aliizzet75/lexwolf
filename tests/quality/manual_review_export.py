"""Anwalt-Review-Export: Generiert DOCX-Exportdatei mit 20 zufälligen Testfällen."""
import json
import random
from pathlib import Path

DATASET_PATH = Path(__file__).parent / "test_dataset.json"
EXPORT_PATH = Path(__file__).parent / "manual_review_export.docx"


def load_cases() -> list:
    """Lädt alle Testfälle aus dem Dataset."""
    try:
        with open(DATASET_PATH) as f:
            return json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        return []


def select_cases(n: int = 20) -> list:
    """Wählt n zufällige Fälle für manuelle Anwalt-Prüfung aus."""
    cases = load_cases()
    if len(cases) <= n:
        return cases
    return random.sample(cases, n)


def generate_export(output_path: str | None = None) -> str:
    """Erstellt DOCX-Exportdatei mit ausgewählten Testfällen für Anwalt-Review.

    Returns: Pfad zur erzeugten Datei
    """
    cases = select_cases(20)
    out = output_path or str(EXPORT_PATH)

    try:
        from docx import Document
        doc = Document()
        doc.add_heading("LexWolf — Anwalt-Review Export", 0)
        doc.add_paragraph(f"Anzahl Fälle: {len(cases)}")
        doc.add_paragraph("Bitte jeden Fall bewerten: Score 1-5 (1=falsch, 5=korrekt)")
        doc.add_paragraph("")

        for i, case in enumerate(cases, 1):
            doc.add_heading(f"Fall {i}: {case.get('frage', '')[:80]}", level=2)
            doc.add_paragraph(f"Erwartete §§: {', '.join(case.get('erwartete_paragraphen', []))}")
            doc.add_paragraph(f"Schwierigkeit: {case.get('schwierigkeit', 'unbekannt')}")
            doc.add_paragraph("Anwalt-Bewertung (1-5): ___")
            doc.add_paragraph("Kommentar: ________________________________")
            doc.add_paragraph("")

        doc.save(out)
    except ImportError:
        # Fallback: einfache Textdatei wenn python-docx nicht verfügbar
        out = out.replace(".docx", ".txt")
        with open(out, "w", encoding="utf-8") as f:
            f.write("LexWolf — Anwalt-Review Export\n")
            f.write(f"Anzahl Fälle: {len(cases)}\n\n")
            for i, case in enumerate(cases, 1):
                f.write(f"Fall {i}: {case.get('frage', '')}\n")
                f.write(f"Erwartete §§: {', '.join(case.get('erwartete_paragraphen', []))}\n")
                f.write("Bewertung (1-5): ___\n\n")

    return out


if __name__ == "__main__":
    path = generate_export()
    print(f"Export erstellt: {path}")
